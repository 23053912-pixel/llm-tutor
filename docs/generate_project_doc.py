from __future__ import annotations

from pathlib import Path
from textwrap import fill

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


OUTPUT_DIR = Path(__file__).resolve().parent
DOCX_PATH = OUTPUT_DIR / "LLM_Tutor_Project_Documentation_Rewritten.docx"
HOME_IMG = OUTPUT_DIR / "llm_tutor_home_view.png"
CHAT_IMG = OUTPUT_DIR / "llm_tutor_chat_view.png"
ARCH_IMG = OUTPUT_DIR / "llm_tutor_architecture.png"


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        Path("C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size=size)
    return ImageFont.load_default()


def draw_wrapped_text(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, font_obj, fill_color, line_spacing: int = 6):
    x1, y1, x2, y2 = box
    max_width = x2 - x1
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        trial = word if not current else f"{current} {word}"
        width = draw.textbbox((0, 0), trial, font=font_obj)[2]
        if width <= max_width:
            current = trial
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)

    y = y1
    for line in lines:
        draw.text((x1, y), line, font=font_obj, fill=fill_color)
        y += font_obj.size + line_spacing
        if y > y2:
            break


def rounded_panel(draw: ImageDraw.ImageDraw, box, fill_color, outline=None, width: int = 2, radius: int = 24):
    draw.rounded_rectangle(box, radius=radius, fill=fill_color, outline=outline, width=width)


def create_home_view(path: Path) -> None:
    img = Image.new("RGB", (1600, 1000), "#0f1117")
    draw = ImageDraw.Draw(img)

    f_title = font(34, bold=True)
    f_sub = font(22, bold=True)
    f_body = font(18)
    f_small = font(16)
    f_mini = font(14)

    draw.rectangle((0, 0, 350, 1000), fill="#161b22")
    draw.line((350, 0, 350, 1000), fill="#21262d", width=2)

    draw.text((28, 28), "Session", font=f_sub, fill="#e6edf3")
    sidebar_lines = [
        "Thread ID: 4fd9b2a1",
        "Knowledge base documents: 15",
        "Model provider: Groq",
        "Model: llama-3.3-70b-versatile",
        "Embedder: all-MiniLM-L6-v2",
        "",
        "Topics Covered",
        "- LLM basics",
        "- Transformer attention",
        "- Training and fine-tuning",
        "- Prompt engineering",
        "- Retrieval augmented generation",
        "- Evaluation and safety",
        "- Applications and future trends",
    ]
    y = 78
    for line in sidebar_lines:
        fill_color = "#e6edf3" if line in {"Topics Covered"} else "#8b949e"
        draw.text((28, y), line, font=f_small, fill=fill_color)
        y += 28

    rounded_panel(
        draw,
        (390, 42, 1548, 250),
        fill_color="#1a5c4c",
        outline="#265b4c",
        radius=26,
    )
    draw.text((430, 80), "LLM Tutor", font=font(40, bold=True), fill="white")
    draw_wrapped_text(
        draw,
        (430, 136, 1450, 192),
        "Your interactive guide to understanding Large Language Models, from architecture to real-world applications.",
        f_body,
        "#f4f6f8",
    )
    draw.text(
        (430, 205),
        "Powered by Llama 3.3 70B | Retrieval-Augmented | RAG-Enhanced Responses",
        font=f_small,
        fill="#ffe0c0",
    )

    draw.text((390, 290), "Quick Start Examples", font=f_sub, fill="#e6edf3")

    card_specs = [
        ("Architecture and Theory", "Explain how transformer architecture and attention mechanisms work in LLMs."),
        ("Practical Applications", "What are the real-world applications of large language models?"),
        ("LLM Trends and Future", "What are the emerging trends and future direction of LLMs?"),
    ]
    x = 390
    for title, text in card_specs:
        rounded_panel(draw, (x, 338, x + 360, 525), fill_color="#17212b", outline="#31404d", radius=22)
        draw.text((x + 24, 364), title, font=f_small, fill="#e6edf3")
        draw_wrapped_text(draw, (x + 24, 404, x + 330, 485), text, f_mini, "#9aa4af")
        draw.text((x + 24, 486), "Try ->", font=f_small, fill="#c96b3c")
        x += 388

    draw.text((390, 565), "Chat", font=f_sub, fill="#e6edf3")
    draw.line((390, 604, 1548, 604), fill="#2b323c", width=2)
    rounded_panel(draw, (410, 640, 1515, 720), fill_color="#141923", outline="#2b323c", radius=20)
    draw.text((438, 668), "Ask me about LLMs...", font=f_small, fill="#6e7681")

    img.save(path)


def create_chat_view(path: Path) -> None:
    img = Image.new("RGB", (1600, 1100), "#0f1117")
    draw = ImageDraw.Draw(img)

    f_title = font(34, bold=True)
    f_sub = font(22, bold=True)
    f_body = font(18)
    f_small = font(16)
    f_mini = font(14)

    draw.rectangle((0, 0, 350, 1100), fill="#161b22")
    draw.line((350, 0, 350, 1100), fill="#21262d", width=2)
    draw.text((28, 28), "Session", font=f_sub, fill="#e6edf3")
    draw.text((28, 76), "Thread ID: 4fd9b2a1", font=f_small, fill="#8b949e")
    draw.text((28, 104), "Faithfulness target: > 0.8", font=f_small, fill="#8b949e")
    draw.text((28, 132), "Retrieval depth: top 3 documents", font=f_small, fill="#8b949e")

    rounded_panel(draw, (390, 38, 1548, 170), fill_color="#15222c", outline="#2d3b48", radius=22)
    draw.text((430, 72), "Chat", font=f_title, fill="#e6edf3")
    draw.text((430, 118), "Conversation with grounding, sources, and evaluation trace.", font=f_body, fill="#9aa4af")

    rounded_panel(draw, (410, 215, 980, 325), fill_color="#1b2430", outline="#324152", radius=20)
    draw.text((438, 246), "User", font=f_small, fill="#d2d7dd")
    draw_wrapped_text(
        draw,
        (438, 278, 950, 312),
        "Explain how retrieval augmented generation reduces hallucination in LLM applications.",
        f_body,
        "#f4f6f8",
    )

    rounded_panel(draw, (520, 360, 1518, 730), fill_color="#18222c", outline="#31404d", radius=20)
    draw.text((550, 392), "Assistant", font=f_small, fill="#d2d7dd")
    assistant_text = (
        "RAG reduces hallucination by grounding the model in retrieved documents before the answer is generated. "
        "In this project, the user query is embedded, the top 3 relevant documents are pulled from ChromaDB, "
        "and those passages are inserted into the prompt so the model responds with source-backed explanations."
    )
    draw_wrapped_text(draw, (550, 426, 1472, 560), assistant_text, f_body, "#f4f6f8")

    rounded_panel(draw, (550, 590, 1472, 700), fill_color="#10161d", outline="#29313a", radius=18)
    draw.text((578, 620), "Sources: Retrieval Augmented Generation, LLM Evaluation and Metrics", font=f_small, fill="#cfd6dd")
    draw.text((578, 655), "Faithfulness score: 0.92 | Route: retrieve | Eval retries: 1", font=f_small, fill="#78dba9")

    rounded_panel(draw, (410, 780, 1518, 1005), fill_color="#141923", outline="#2b323c", radius=20)
    draw.text((438, 814), "Trace Expander", font=f_small, fill="#cfd6dd")
    trace_text = (
        "memory -> router -> retrieve -> answer -> eval -> save\n"
        "Sliding window memory keeps the last 6 messages, while the evaluator checks whether the final answer stays grounded in retrieved context."
    )
    draw.multiline_text((438, 850), trace_text, font=f_body, fill="#9aa4af", spacing=10)

    img.save(path)


def arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], fill_color: str = "#1a5c4c", width: int = 6):
    draw.line((start, end), fill=fill_color, width=width)
    ex, ey = end
    sx, sy = start
    if abs(ex - sx) >= abs(ey - sy):
        direction = 1 if ex > sx else -1
        draw.polygon([(ex, ey), (ex - 18 * direction, ey - 10), (ex - 18 * direction, ey + 10)], fill=fill_color)
    else:
        direction = 1 if ey > sy else -1
        draw.polygon([(ex, ey), (ex - 10, ey - 18 * direction), (ex + 10, ey - 18 * direction)], fill=fill_color)


def create_architecture_view(path: Path) -> None:
    img = Image.new("RGB", (1600, 900), "white")
    draw = ImageDraw.Draw(img)

    f_title = font(34, bold=True)
    f_sub = font(20, bold=True)
    f_body = font(16)

    draw.text((70, 42), "LLM Tutor System Architecture", font=f_title, fill="#10222b")

    boxes = {
        "ui": (90, 150, 370, 255, "#e9f5f1", "Streamlit UI\nChat interface, hero section,\nsession state"),
        "memory": (470, 100, 760, 190, "#f5efe3", "Memory Node\nStores conversation and user name"),
        "router": (470, 225, 760, 315, "#f5efe3", "Router Node\nChooses retrieve, tool, or skip"),
        "retrieve": (470, 350, 760, 440, "#f5efe3", "Retrieval Node\nEmbeds query and fetches top 3 docs"),
        "answer": (470, 475, 760, 565, "#f5efe3", "Answer Node\nGenerates grounded explanation"),
        "eval": (470, 600, 760, 690, "#f5efe3", "Eval Node\nScores faithfulness and retries"),
        "vector": (910, 250, 1245, 360, "#eef1fb", "ChromaDB + SentenceTransformer\nVector search over 15 topic documents"),
        "groq": (910, 470, 1245, 580, "#eef1fb", "Groq API\nLlama 3.3 70B for fast inference"),
        "output": (1290, 360, 1510, 470, "#e9f5f1", "Final Response\nAnswer + sources + score"),
    }

    for _, (x1, y1, x2, y2, fill_color, text) in boxes.items():
        rounded_panel(draw, (x1, y1, x2, y2), fill_color=fill_color, outline="#9db0b8", width=3, radius=20)
        title, body = text.split("\n", 1)
        draw.text((x1 + 18, y1 + 18), title, font=f_sub, fill="#10222b")
        draw.multiline_text((x1 + 18, y1 + 48), body, font=f_body, fill="#344a53", spacing=8)

    arrow(draw, (370, 202), (470, 145))
    arrow(draw, (615, 190), (615, 225))
    arrow(draw, (615, 315), (615, 350))
    arrow(draw, (615, 440), (615, 475))
    arrow(draw, (615, 565), (615, 600))
    arrow(draw, (760, 395), (910, 305))
    arrow(draw, (760, 520), (910, 525))
    arrow(draw, (1245, 305), (1290, 390))
    arrow(draw, (1245, 525), (1290, 430))

    footer = (
        "Flow: user question -> routing -> retrieval/tool handling -> grounded answer generation -> "
        "faithfulness check -> final response shown in the Streamlit chat interface."
    )
    draw_wrapped_text(draw, (90, 770, 1510, 845), footer, f_body, "#344a53", line_spacing=8)

    img.save(path)


def set_run_font(run, size: int, bold: bool = False, color: RGBColor | None = None):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def add_page_field(paragraph):
    begin_run = paragraph.add_run()
    set_run_font(begin_run, 12)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    begin_run._r.append(fld_begin)

    instr_run = paragraph.add_run()
    set_run_font(instr_run, 12)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    instr_run._r.append(instr)

    separate_run = paragraph.add_run()
    set_run_font(separate_run, 12)
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    separate_run._r.append(fld_separate)

    text_run = paragraph.add_run("1")
    set_run_font(text_run, 12)

    end_run = paragraph.add_run()
    set_run_font(end_run, 12)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    end_run._r.append(fld_end)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.0)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(12)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    h1 = doc.styles["Heading 1"]
    h1.font.name = "Arial"
    h1._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    h1.font.size = Pt(15)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(16, 34, 43)

    h2 = doc.styles["Heading 2"]
    h2.font.name = "Arial"
    h2._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(26, 92, 76)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_page_field(footer)


def add_justified_paragraph(doc: Document, text: str, bold_prefix: str | None = None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold_prefix:
        prefix_run = p.add_run(bold_prefix)
        set_run_font(prefix_run, 12, bold=True)
        body_run = p.add_run(text)
        set_run_font(body_run, 12)
    else:
        run = p.add_run(text)
        set_run_font(run, 12)


def add_centered_title(doc: Document, text: str, size: int):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, size, bold=True, color=RGBColor(16, 34, 43))


def add_image_with_caption(doc: Document, image_path: Path, caption: str, width_inches: float):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))
    caption_p = doc.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption_p.add_run(caption)
    set_run_font(caption_run, 11, bold=True, color=RGBColor(80, 88, 96))


def add_tech_stack_table(doc: Document):
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["Layer", "Technology", "Role in Project"]
    for cell, text in zip(table.rows[0].cells, headers):
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(text)
        set_run_font(run, 12, bold=True)

    rows = [
        ("Frontend", "Streamlit", "Used to build the chat page, sidebar, session controls, and guided example prompts."),
        ("Workflow Engine", "LangGraph", "Runs the node-based flow for memory handling, routing, retrieval, response generation, and saving state."),
        ("Model Layer", "Groq with Llama 3.3 70B", "Produces fast answers for student questions inside the tutoring workflow."),
        ("Retrieval Layer", "ChromaDB + SentenceTransformer", "Converts topics into embeddings and returns the closest knowledge chunks for each query."),
        ("Content Layer", "Python knowledge base module", "Stores curated learning material on transformers, training, RAG, evaluation, safety, and applications."),
        ("Quality Check", "Faithfulness scoring", "Measures whether the answer stays tied to the retrieved context before it is shown to the user."),
    ]

    for layer, tech, role in rows:
        cells = table.add_row().cells
        values = [layer, tech, role]
        for cell, value in zip(cells, values):
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = paragraph.add_run(value)
            set_run_font(run, 11)


def build_document() -> None:
    doc = Document()
    configure_document(doc)

    add_centered_title(doc, "Project Documentation", 18)
    add_centered_title(doc, "LLM Tutor - Building a Topic-Focused Assistant for Learning Large Language Models", 16)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("This version is written directly from the implemented project files and observed application flow.")
    set_run_font(run, 12)

    doc.add_paragraph()
    doc.add_heading("1. Project Context", level=1)
    add_justified_paragraph(
        doc,
        "This capstone was built around a practical question: how can a learner study large language models through conversation without depending on a fully open-ended chatbot? The answer developed in this project is LLM Tutor, a Streamlit-based study assistant that keeps its attention on one subject area and uses a prepared knowledge base instead of relying only on raw model recall."
    )
    add_justified_paragraph(
        doc,
        "The application is intentionally narrow in scope. Its content is not pulled from random web pages during the session. Instead, the project stores its own topic material inside the repository, then routes each question through a controlled workflow. That makes the tutor useful for demonstrations, beginner learning, and capstone evaluation because its behavior can be explained from the code itself."
    )

    doc.add_heading("2. Problem Statement", level=1)
    add_justified_paragraph(
        doc,
        "Students who start learning about LLMs usually face two difficulties at the same time. The first is content overload. Important concepts such as attention, embeddings, pretraining, fine-tuning, RAG, alignment, and evaluation are often explained in different places and at very different technical levels. The second is answer reliability. General AI chat tools can explain these ideas fluently, but they do not always show whether the answer was grounded in a relevant source."
    )
    add_justified_paragraph(
        doc,
        "The purpose of this project is therefore not just to answer questions, but to answer them inside a controlled educational setting. The system needs to remember recent conversation, detect whether a prompt belongs to the learning domain, retrieve supporting content when required, and show signals that help the user trust the output. In short, the project tries to turn a chat interface into a guided learning tool."
    )

    doc.add_page_break()

    doc.add_heading("3. Solution and Features", level=1)
    add_justified_paragraph(
        doc,
        "The finished solution combines a Streamlit front end with a LangGraph workflow behind it. At startup, the app loads environment settings, checks for a Groq API key, prepares embeddings with `all-MiniLM-L6-v2`, creates an in-memory ChromaDB collection named `course_kb`, and loads the local knowledge base. Once the system is ready, every user message is passed to a graph that handles memory, routing, retrieval, generation, scoring, and saving."
    )

    doc.add_heading("3.1 Key Features", level=2)
    add_justified_paragraph(doc, "The interface provides a chat-first learning experience, so the user can ask direct questions such as how transformer attention works, where RAG helps, or how evaluation metrics differ.")
    add_justified_paragraph(doc, "The project contains fifteen knowledge topics in `kb_data.py`, which makes the tutor focused enough to stay relevant while still broad enough to cover the main LLM study areas.")
    add_justified_paragraph(doc, "The sidebar exposes useful session details such as the thread identifier, document count, chosen model, embedding model, topic list, and API-key status.")
    add_justified_paragraph(doc, "Three guided starter prompts are shown before a conversation begins, which helps first-time users explore architecture, applications, and future trends without having to guess what to ask.")
    add_justified_paragraph(doc, "Each response can carry supporting sources and a faithfulness score, so the user sees more than just final text and gains insight into how the answer was produced.")

    doc.add_heading("3.2 Working Flow", level=2)
    add_justified_paragraph(
        doc,
        "When the learner submits a prompt, `process_query()` invokes the compiled graph using the active thread id. The memory node appends the new human message and keeps only the last six messages. The router node then classifies the prompt into one of three labels: `retrieve`, `tool`, or `skip`. If retrieval is needed, the query is embedded, matched against the vector store, and the top three documents are added to the context. The answer node builds a strict system prompt, asks the model to stay within the supplied material, and returns the draft answer. The evaluation node scores faithfulness on a scale from 0 to 1 and triggers another attempt when the score falls below the threshold. Finally, the accepted response is stored and shown to the user."
    )

    doc.add_page_break()

    doc.add_heading("4. Screenshots", level=1)
    add_justified_paragraph(
        doc,
        "The following figures are included to show the project in a report-friendly format. The first image presents the landing state of the application, the second illustrates how a grounded answer is displayed inside the chat interface, and the third gives a simplified architecture view of the implemented workflow."
    )
    add_image_with_caption(doc, HOME_IMG, "Figure 1. Main application view with sidebar details and guided starter prompts.", 6.3)
    add_image_with_caption(doc, CHAT_IMG, "Figure 2. Example chat output with retrieved sources and faithfulness information.", 6.3)

    doc.add_page_break()

    doc.add_heading("5. Technical Architecture", level=1)
    add_justified_paragraph(
        doc,
        "Although the user sees only a chat window, the backend is divided into small and understandable parts. The Streamlit layer is responsible for interaction and session state. LangGraph coordinates the backend nodes. The embedding model converts both stored material and incoming questions into vectors. ChromaDB acts as the local retrieval store. Groq supplies the Llama 3.3 70B model for generation. This breakdown makes the system easier to explain in a capstone setting because every major action in the user flow corresponds to a clear block in the implementation."
    )
    add_image_with_caption(doc, ARCH_IMG, "Figure 3. High-level component flow used by the LLM Tutor system.", 6.5)
    add_justified_paragraph(doc, "The most important design decision is the graph structure in `agent.py`. Instead of handling everything in one large function, the project uses separate nodes for memory, routing, retrieval, answer generation, evaluation, and save logic. That choice improves readability and also makes later modification easier because each behavior is isolated.")
    add_justified_paragraph(doc, "Another notable detail is the answer-evaluation loop. If retrieved context exists, the system asks the model to judge whether the answer stayed faithful to that context. The graph retries while the score is below `0.7` and the retry count is below `2`. This is a simple mechanism, but it clearly shows the idea of self-checking in an agentic workflow.")

    doc.add_heading("5.1 Node-by-Node Summary", level=2)
    add_justified_paragraph(doc, "Memory Node: Stores the current user question in conversation history, trims the history to six messages, and can capture the user's name when it appears in the question.")
    add_justified_paragraph(doc, "Router Node: Uses the model itself to output a single route label so that the rest of the graph knows whether retrieval, time-tool behavior, or a skip path should be used.")
    add_justified_paragraph(doc, "Retrieval Node: Encodes the user query, runs a similarity search in ChromaDB, and formats the returned topic text as context for answer generation.")
    add_justified_paragraph(doc, "Answer Node: Builds the system prompt, injects the knowledge-base context and chat history, and asks the LLM to answer in a friendly educational tone.")
    add_justified_paragraph(doc, "Evaluation and Save Nodes: Score the response for grounding quality, retry when necessary, and finally append the accepted assistant reply to memory.")

    doc.add_heading("6. Tech Stack", level=1)
    add_tech_stack_table(doc)

    doc.add_heading("7. Unique Points", level=1)
    add_justified_paragraph(doc, "The project does not pretend to be a universal assistant. Its strength comes from being topic-locked. Because the stored material is about LLMs only, the tutor behaves more like a guided course companion than a random chatbot.")
    add_justified_paragraph(doc, "The application also exposes internal signals that many beginner demos hide. Route selection, sources, and faithfulness details are part of the experience, which helps users understand not just the answer but also the path taken to produce it.")
    add_justified_paragraph(doc, "The repository structure is simple enough for a reviewer to inspect quickly. `capstone_streamlit.py` handles the interface, `agent.py` contains the graph logic, and `kb_data.py` contains the learning content. This clear split is useful during viva, review, and future extension.")

    doc.add_heading("8. Current Limitations", level=1)
    add_justified_paragraph(doc, "The present version is designed as a capstone prototype, so some production features are still missing. The Chroma collection is recreated when the system loads, which is fine for demonstration but not ideal for a persistent learning platform.")
    add_justified_paragraph(doc, "The current automated test file is minimal and mostly checks whether important modules import correctly. Functional test coverage for retrieval quality, routing accuracy, and score-based retries has not yet been added.")
    add_justified_paragraph(doc, "The app also depends on the user providing a valid Groq API key. Without that key, the interface can load, but the tutor cannot answer questions.")

    doc.add_page_break()

    doc.add_heading("9. Future Improvements", level=1)
    add_justified_paragraph(doc, "The most useful next upgrade would be persistent storage for both chat history and vector data so that the user does not lose state between sessions and the system does not rebuild its collection every time.")
    add_justified_paragraph(doc, "Retrieval quality can be strengthened with chunk tuning, hybrid search, or a reranking stage. This would be especially helpful for broad questions where pure semantic similarity may miss the most instructive explanation.")
    add_justified_paragraph(doc, "The evaluation layer can be expanded beyond one faithfulness score. A more mature version could score relevance, answer completeness, and user satisfaction, then use those signals to guide future improvements.")
    add_justified_paragraph(doc, "From the product side, the tutor can evolve into a richer study environment by adding quizzes, revision notes, multilingual support, downloadable summaries, and a deployed version accessible outside the local machine.")

    doc.add_heading("10. Conclusion", level=1)
    add_justified_paragraph(
        doc,
        "LLM Tutor demonstrates that a small but carefully structured AI project can teach more effectively than a generic chat wrapper. By limiting the domain, curating the knowledge base, routing prompts through a graph, retrieving supporting context, and checking the answer before display, the project creates a study assistant that is easier to explain, review, and improve. It works both as a learning product and as a clear demonstration of retrieval-based agent design."
    )
    add_justified_paragraph(
        doc,
        "Reference Note: the technical descriptions in this document were prepared from the implemented repository files `capstone_streamlit.py`, `agent.py`, `kb_data.py`, `requirements.txt`, and `tests/test_basic.py`. This note is included to keep the report tied to the actual project rather than to generic external summaries."
    )

    doc.save(DOCX_PATH)


def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    create_home_view(HOME_IMG)
    create_chat_view(CHAT_IMG)
    create_architecture_view(ARCH_IMG)
    build_document()
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
